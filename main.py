#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import tkinter as tk
from tkinter import filedialog, messagebox
import fitz  # PyMuPDF
from PIL import Image
import io
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# 创建保存目录
save_dir = "output"


def main():
    # 创建Tk根窗口并隐藏
    root = tk.Tk()
    root.withdraw()

    # 打开文件选择对话框
    file_paths = filedialog.askopenfilenames(
        title="选择PDF文件", filetypes=[("PDF files", "*.pdf")]
    )

    if file_paths:
        # 检查保存目录是否存在文件
        if os.path.exists(save_dir) and os.listdir(save_dir):
            # 询问是否清空文件夹
            result = messagebox.askyesno("确认", "是否清空保存文件夹？")
            if result:
                for filename in os.listdir(save_dir):
                    file_path = os.path.join(save_dir, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.unlink(file_path)
                    except Exception as e:
                        print(f"删除文件失败: {e}")

        # 创建保存目录（如果不存在）
        os.makedirs(save_dir, exist_ok=True)

        # 创建Word文档
        doc = Document()

        # 设置页边距为1.27厘米
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

        # 打开保存文件夹
        try:
            os.startfile(save_dir)  # Windows
        except:
            try:
                os.system(f'open "{save_dir}"')  # macOS
            except:
                os.system(f'xdg-open "{save_dir}"')  # Linux

        # 临时存储转换后的图片
        image_paths = []

        # 转换每个PDF文件
        for i, file_path in enumerate(file_paths):
            try:
                # 获取不带扩展名的文件名
                file_name = os.path.splitext(os.path.basename(file_path))[0]

                # 打开PDF文件
                pdf_document = fitz.open(file_path)

                # 获取第一页
                page = pdf_document[0]

                # 将页面渲染为图像
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom
                pix = page.get_pixmap(matrix=mat)

                # 保存图像到临时文件
                temp_img_path = os.path.join(save_dir, f"temp_{i}_{file_name}.png")
                pix.save(temp_img_path)
                image_paths.append(temp_img_path)
                print(f"已保存临时图片: {temp_img_path}")

                # 关闭PDF文档
                pdf_document.close()

            except Exception as e:
                print(f"处理文件 {file_path} 时出错: {e}")

        # 将图片添加到Word文档，每页两张
        for i in range(0, len(image_paths), 2):
            # 添加新段落
            if i > 0:  # 除了第一页，其他页需要添加分页符
                doc.add_page_break()

            # 添加第一张图片
            if i < len(image_paths):
                paragraph1 = doc.add_paragraph()
                paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run1 = paragraph1.add_run()
                run1.add_picture(image_paths[i], width=Cm(18))

            # 添加第二张图片（如果存在）
            if i + 1 < len(image_paths):
                paragraph2 = doc.add_paragraph()
                paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run2 = paragraph2.add_run()
                run2.add_picture(image_paths[i + 1], width=Cm(18))

        # 生成带时间戳的文件名
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        doc_path = os.path.join(save_dir, f"output_{timestamp}.docx")

        # 保存Word文档
        doc.save(doc_path)
        print(f"已创建Word文档: {doc_path}")

        # 删除临时图片
        for img_path in image_paths:
            try:
                os.remove(img_path)
            except Exception as e:
                print(f"删除临时文件 {img_path} 时出错: {e}")


if __name__ == "__main__":
    main()
